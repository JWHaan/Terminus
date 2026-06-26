from __future__ import annotations

import io
import re
from pathlib import Path
from typing import BinaryIO

from dateutil import parser as date_parser
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_text(uploaded_file: BinaryIO, filename: str) -> str:
    """
    Extract text from PDF, DOCX, TXT, or MD files uploaded through Streamlit.
    """
    suffix = Path(filename).suffix.lower()
    raw = uploaded_file.read()

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(raw))
        pages: list[str] = []

        for i, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            pages.append(f"\n--- Page {i} ---\n{page_text}")

        return "\n".join(pages).strip()

    if suffix == ".docx":
        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip()

    if suffix in {".txt", ".md"}:
        return raw.decode("utf-8", errors="ignore").strip()

    raise ValueError(f"Unsupported file type: {suffix}. Upload PDF, DOCX, TXT, or MD.")


def simple_deadline_fallback(text: str, filename: str) -> list[dict]:
    """
    Basic no-API parser.
    It searches for date-like lines near deadline/exam keywords.
    """
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    keywords = re.compile(
        r"\b(deadline|due|submission|submit|exam|test|quiz|presentation|project|assignment|lab)\b",
        re.I,
    )

    date_like = re.compile(
        r"(\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|"
        r"\b\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{2,4}\b|"
        r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{2,4}\b)",
        re.I,
    )

    results: list[dict] = []

    for idx, line in enumerate(lines):
        window = " ".join(lines[max(0, idx - 1) : min(len(lines), idx + 2)])

        if not keywords.search(window):
            continue

        match = date_like.search(window)

        if not match:
            continue

        try:
            parsed = date_parser.parse(
                match.group(0),
                dayfirst=True,
                fuzzy=True,
            ).date()

        except Exception:
            continue

        if re.search(r"\bexam|test|quiz\b", window, re.I):
            kind = "exam"
        elif re.search(r"\bpresentation\b", window, re.I):
            kind = "presentation"
        elif re.search(r"\bproject\b", window, re.I):
            kind = "project"
        elif re.search(r"\blab\b", window, re.I):
            kind = "lab"
        else:
            kind = "assignment"

        title = line[:100] or f"Deadline from {filename}"

        results.append(
            {
                "course_code": None,
                "course_name": None,
                "title": title,
                "kind": kind,
                "due_date": parsed.isoformat(),
                "due_time": None,
                "weight": None,
                "description": window[:500],
                "source_quote": line[:180],
                "total_estimated_hours": 6.0 if kind != "exam" else 12.0,
                "checklist": default_checklist(kind),
            }
        )

    return results


def default_checklist(kind: str) -> list[dict]:
    if kind == "exam":
        return [
            {
                "title": "Collect tested topics",
                "details": "List topics, formulas, chapters, and tutorial questions tested.",
                "estimated_hours": 1.0,
                "priority": "high",
            },
            {
                "title": "Make summary notes",
                "details": "Condense the main concepts into short revision notes.",
                "estimated_hours": 3.0,
                "priority": "high",
            },
            {
                "title": "Practice questions",
                "details": "Attempt tutorial questions, practice papers, and similar problems.",
                "estimated_hours": 5.0,
                "priority": "high",
            },
            {
                "title": "Fix weak spots",
                "details": "Redo mistakes and write down why the correct method works.",
                "estimated_hours": 2.0,
                "priority": "medium",
            },
            {
                "title": "Final review",
                "details": "Do light revision and prepare materials needed for the exam.",
                "estimated_hours": 1.0,
                "priority": "medium",
            },
        ]

    return [
        {
            "title": "Read instructions",
            "details": "Highlight deliverables, marking criteria, and submission format.",
            "estimated_hours": 0.5,
            "priority": "high",
        },
        {
            "title": "Plan structure",
            "details": "Break the work into sections and decide what evidence, code, or data is needed.",
            "estimated_hours": 1.0,
            "priority": "high",
        },
        {
            "title": "First draft / build",
            "details": "Create the main submission content before polishing.",
            "estimated_hours": 3.0,
            "priority": "high",
        },
        {
            "title": "Improve and check",
            "details": "Fix gaps, add references, test code, or review calculations.",
            "estimated_hours": 1.0,
            "priority": "medium",
        },
        {
            "title": "Final submission",
            "details": "Export/upload the correct file and confirm submission receipt.",
            "estimated_hours": 0.5,
            "priority": "high",
        },
    ]